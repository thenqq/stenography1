import docx
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

def take_text(path):
    qwe = []
    for p in path.paragraphs:
        qwe.append(p.text)
        print(p.text)
        p.clear()
    print(qwe)
    return qwe

def split_text(bits, text):
    tre = []
    tline = ''
    for i in range(len(bits)):
        if i != len(bits) - 1:
            if bits[i] == '0' and bits[i+1] == '0':
                tline += text[i]
            elif bits[i] == '1' and bits[i+1] == '1':
                tline += text[i]
            elif bits[i] == '0' and bits[i + 1] == '1':
                tline += text[i]
                tre.append(tline)
                tline = ''
            elif bits[i] == '1' and bits[i + 1] == '0':
                tline += text[i]
                tre.append(tline)
                tline = ''
        else:
            if bits[i] == '0' and bits[i - 1] == '0':
                tline += text[i]
                tre.append(tline)
            elif bits[i] == '1' and bits[i - 1] == '1':
                tline += text[i]
                tre.append(tline)
            elif bits[i] == '0' and bits[i - 1] == '1':
                tline = text[i]
                tre.append(tline)
            elif bits[i] == '1' and bits[i - 1] == '0':
                tline = text[i]
                tre.append(tline)

    return tre

my_text = 'Бог труды любит.'
my_encoding = 'koi8-r'
bit_line = ''.join(format(byte, '08b') for byte in my_text.encode(my_encoding))

print(f'text = {my_text}')
print(f'bits = {bit_line}')
doc = docx.Document('./cont/2.docx')

array_text = take_text(doc)

array_text_split_to_p =[]
for i in array_text:
    if len(i.split('\n')) > 1:
        array_text_split_to_p.append(i.split('\n'))
    else:
        array_text_split_to_p.append(i)

array_text_lengths = []
for i in array_text:
    rwqe = []
    for j in i.split('\n'):
        rwqe.append(len(j))
    array_text_lengths.append([len('|'.join(i.split('\n'))), len(i.split('\n')) - 1, rwqe])

need_paragraphs = 0
for i in range(len(array_text_lengths)):
    need_paragraphs += array_text_lengths[i][0] - array_text_lengths[i][1]
    if need_paragraphs >= len(bit_line):
        need_paragraphs = i + 1
        break

bit_line_split_to_p = []
q = 0
right = 0

for i in range(need_paragraphs):
    if q == 0:
        left = 0
        right += array_text_lengths[i][0]
        bit_line_split_to_p.append(bit_line[left:right])
        q = right
    elif q != 0:
        left = q
        right += array_text_lengths[i][0]

        temp_l = bit_line[left:right]
        if len(temp_l) != 0:
            w = 0
            temp_a = []
            right_2 = 0
            for k in range(len(array_text_lengths[i][2])):
                left_2 = w
                right_2 += array_text_lengths[i][2][k]
                if len(temp_l[left_2:right_2]) == 0:
                    break
                temp_a.append(temp_l[left_2:right_2])
                w = right_2
            bit_line_split_to_p.append(temp_a)
        else:
            temp_l = bit_line[left:]

            qcount = 0
            for ee in range(len(array_text_lengths[i][2])):
                qcount += array_text_lengths[i][2][ee]
                if qcount >= len(temp_l):
                    qcount = ee + 1
                    break

            temp_a = []
            w = 0
            right_2 = 0
            for ee in range(qcount):
                left_2 = w
                right_2 += array_text_lengths[i][2][ee]
                if len(temp_l[left_2:right_2]) != 0:
                    temp_a.append(temp_l[left_2:right_2])
                else:
                    temp_a.append(temp_l[left_2:])
                w = right_2

            bit_line_split_to_p.append(temp_a)

        q = right
print(bit_line_split_to_p)
print(array_text_split_to_p)
print()
for i in range(len(array_text)):
    p = doc.paragraphs[i]
    if i == 0:
        temp_array = split_text(bit_line_split_to_p[i], array_text_split_to_p[i])
        print(temp_array)
        for j in temp_array:
            p.add_run(j)
    elif i < need_paragraphs and i > 0:
        # ['010110011111100011100', '10000011010100110100', '10110101011100010', '01101100100100000110']
        # ['Шла вчера я за водою,', 'А у нас ведро худое.', 'Из-за этого ведра', 'Я наплакалась вчера.']
        if len(bit_line_split_to_p[i]) == len(array_text_split_to_p[i]):
            for j in range(len(bit_line_split_to_p[i])):
                giga_array = split_text(bit_line_split_to_p[i][j], array_text_split_to_p[i][j])
                #print(giga_array)
                for k in giga_array:
                    p.add_run(k)
                if j != len(bit_line_split_to_p[i])-1:
                    p.add_run('\n')

        # ['0011000000110000', '10110010011101010000', '101110']
        # ['Залепила я дыру.', 'Только воду наберу —', 'А она опять наружу', 'Так и льется по ведру.']
        elif len(bit_line_split_to_p[i]) != len(array_text_split_to_p[i]):
            count1 = 0
            count2 = len(array_text_split_to_p[i])
            for j in range(len(bit_line_split_to_p[i])):
                if len(bit_line_split_to_p[i][j]) == len(array_text_split_to_p[i][j]):
                    giga_array = split_text(bit_line_split_to_p[i][j], array_text_split_to_p[i][j])
                    #print(giga_array)
                    for k in giga_array:
                        p.add_run(k)
                    if j != len(bit_line_split_to_p[i]) - 1:
                        p.add_run('\n')
                    count1 += 1
                elif len(bit_line_split_to_p[i][j]) != len(array_text_split_to_p[i][j]):
                    bits_temp = bit_line_split_to_p[i][j]
                    text_temp = array_text_split_to_p[i][j][:len(bits_temp)]
                    giga_array = split_text(bits_temp, text_temp)
                    ost_tt = array_text_split_to_p[i][j][len(bits_temp):]
                    #print(giga_array)
                    for k in giga_array:
                        p.add_run(k)
                    p.add_run(ost_tt)
                    p.add_run('\n')
                    count1 += 1
            revq = array_text_split_to_p[i][::-1]
            for i in range(count2 - count1):
                p.add_run(revq[i])
                if i != count2 - count1 - 1:
                    p.add_run('\n')

    elif i >= need_paragraphs:
        for j in range(len(array_text_split_to_p[i])):
            p.add_run(array_text_split_to_p[i][j])
            if j != len(array_text_split_to_p[i]) - 1:
                p.add_run('\n')

counter = 1
scale_pt = 99
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if i == 0:
        for run in p.runs:
            run.font.size = Pt(22)
            run.font.name = 'Book antiqua'
            if counter % 2 == 1:
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            counter += 1
    elif i != 0:
        for run in p.runs:
            if run.text == '\n':
                continue
            else:
                run.font.size = Pt(14)
                run.font.name = 'Book antiqua'
                if counter % 2 == 1:
                    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                counter += 1

doc.save('./cont/test2.docx')
print('saved to test2')
