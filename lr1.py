
import docx

table_manage = {
    "01000": "return",
    "00010": "slash_n",
    "11111": "eng",
    "11011": "nums",
    "00100": "space",
    "00000": "rus"
}
any_alphabet = {
    "00011": ['A', 'А', '-'],
    "11001": ['B', 'Б', '?'],
    "01110": ['C', 'Ц', ':'],
    "01001": ['D', 'Д', 'Кто там?'],
    "00001": ['E', 'Е', '3'],
    "01101": ['F', 'Ф', 'Э'],
    "11010": ['G', 'Г', 'Ш'],
    "10100": ['H', 'Х', 'Щ'],
    "00110": ['I', 'И', '8'],
    "01011": ['J', 'Й', 'Ю'],
    "01111": ['K', 'К', '('],
    "10010": ['L', 'Л', ')'],
    "11100": ['M', 'М', '.'],
    "01100": ['N', 'Н', ','],
    "11000": ['O', 'О', '9'],
    "10110": ['P', 'П', '0'],
    "10111": ['Q', 'Я', '1'],
    "01010": ['R', 'Р', '4'],
    "00101": ['S', 'С', "'"],
    "10000": ['T', 'Т', '5'],
    "00111": ['U', 'У', '7'],
    "11110": ['V', 'Ж', '='],
    "10011": ['W', 'В', '2'],
    "11101": ['X', 'Ь', '/'],
    "10101": ['Y', 'Ы', '6'],
    "10001": ['Z', 'З', '+']
}

def code_baudot(array):
    strange_line = ''
    flag = 0
    for i in array:
        if i in table_manage:
            manager = table_manage[i]
            if manager == 'eng':
                flag = 0
            elif manager == 'rus':
                flag = 1
            elif manager == 'nums':
                flag = 2
            elif manager == '(return)':
                strange_line += '(CR)'
            elif manager == '(slash_n)':
                strange_line += '(SN)'
            elif manager == '(space)':
                strange_line += ' '

        elif i in any_alphabet:
            strange_line += ''.join(any_alphabet[i][flag])

    return strange_line

def check_color(path):
    for p in docx.Document(path).paragraphs:
        for run in p.runs:
            xmle = run._element
            if 'w:color' in xmle.xml:
                # Извлекаем информацию о цвете шрифта
                color_info = xmle.find('.//w:color', namespaces=xmle.nsmap)
                if color_info is not None:
                    return color_info.attrib, 'w:color'

    return None, None
# w:highlight
def check_highlight(path):
    for p in docx.Document(path).paragraphs:
        for run in p.runs:
            xmle = run._element
            if 'w:highlight' in xmle.xml:
                # Извлекаем информацию о фоне символов
                highlight_info = xmle.find('.//w:highlight', namespaces=xmle.nsmap)
                if highlight_info is not None:
                    return highlight_info.attrib, 'w:highlight'

    return None, None
# w:sz
def check_fsize(path):
    for p in docx.Document(path).paragraphs:
        for run in p.runs:
            xmle = run._element
            if 'w:sz' in xmle.xml:
                # Извлекаем информацию о размере шрифта
                fsize_info = xmle.find('.//w:sz', namespaces=xmle.nsmap)
                if fsize_info is not None:
                    return fsize_info.attrib, 'w:sz'

    return None, None
# w:w
def check_scale(path):
    for p in docx.Document(path).paragraphs:
        for run in p.runs:
            xmle = run._element
            if 'w:w' in xmle.xml:
                scale_info = xmle.find('.//w:w', namespaces=xmle.nsmap)
                if scale_info is not None:
                    return scale_info.attrib, 'w:w'

    return None, None

def check_space(path):
    for p in docx.Document(path).paragraphs:
        for run in p.runs:
            xmle = run._element
            if 'w:spacing' in xmle.xml:
                space_info = xmle.find('.//w:spacing', namespaces=xmle.nsmap)
                if space_info is not None:
                    return space_info.attrib, 'w:spacing'

    return None, None

def find_parameter(path):
    parameter = ''

    f_color, color_flag = check_color(path)
    try:
        q = str(f_color).split(':')[2][1:-1]
    except:
        1
    if (f_color is not None) and (q != "'000000'"):
        print(f"font color: {q}")
    elif (f_color is None) or (q == "'000000'"):
        print("no changes f color")
        color_flag = None

    # Информация о цвете бэка
    highlight, highlight_flag = check_highlight(path)
    try:
        q = str(highlight).split(':')[2][1:-1]
    except:
        1
    if (highlight is not None) and (q != "'white'"):
        print(f"hl color: {q}")
    elif (highlight is None) or (q == "'white'"):
        print(f"no changes hl color")
        highlight_flag = None

    f_size, size_flag = check_fsize(path)
    try:
        q = str(f_size).split(':')[2][1:-1]
    except:
        1
    if (f_size is not None) and ((q != "'24'") or (q != "'28'")):
        print(f"font size: {q}")
    elif (f_size is None) or (q == "'24'") or (q == "'28'"):
        print("no changes f size")
        size_flag = None

    scale, scale_flag = check_scale(path)
    try:
        q = str(scale).split(':')[2][:-1]
    except:
        1
    if (scale is not None):
        print(f"font scale: {q}")
    elif (scale is None):
        print("no changes f scale")
        scale_flag = None

    character_spacing, space_flag = check_space(path)
    try:
        q = str(character_spacing).split(':')[2][:-1]
    except:
        1
    if character_spacing is not None:
        print(f"font spacing: {q}")
    elif character_spacing is None:
        print("no changes f spacing")
        space_flag = None

    print('')

    flags_array = [color_flag, highlight_flag, size_flag, scale_flag, space_flag]
    for i in flags_array:
        if i:
            parameter = i
            break

    return parameter, flags_array

def zeros_ones(path, parameter, flags):
    giga_line = ''

    for p in docx.Document(path).paragraphs:
        for i in range(len(p.runs)):
            print(f'our element(text): {p.runs[i].text}')
            if parameter in p.runs[i]._element.xml:
                char_info = p.runs[i]._element.find(f'.//{parameter}', namespaces=p.runs[i]._element.nsmap)
                if char_info is not None:
                    char_info_q = char_info.attrib
                    char_info_qw = str(char_info_q).split(':')[2][2:-2]
                    print(f'{parameter} value is {char_info_qw}')
                    # print(len(str(paragraph.runs[i].text)))

                    if char_info_qw != '000000' and flags[0]:
                        giga_line += '1' * len(str(p.runs[i].text))
                        print('has changed')
                    elif char_info_qw == '000000' and flags[0]:
                        giga_line += '0' * len(str(p.runs[i].text))
                        print('hasnt change')

                    if char_info_qw != 'white' and flags[1]:
                        giga_line += '1' * len(str(p.runs[i].text))
                        print('has changed')

                    if char_info_qw != '24' and flags[2]:
                        giga_line += '1' * len(str(p.runs[i].text))
                        print('has changed')

                    elif char_info_qw == '24' and flags[2]:
                        giga_line += '0' * len(str(p.runs[i].text))
                        print('hasnt changed')

                    if char_info_qw != '100' and flags[3]:
                        giga_line += '1' * len(str(p.runs[i].text))
                        print('has changed')

                    if char_info_qw == '20' and flags[4]:
                        giga_line += '1' * len(str(p.runs[i].text))
                        print('has changed')
                    print()
            elif parameter not in p.runs[i]._element.xml:
                print('hasnt changed')
                giga_line += '0' * len(str(p.runs[i].text))

    return giga_line

def decode(line):
    line_by_five = []
    for i in range(len(line)):
        try:
            line[i + 5]
            temp_line = ''
            for j in range(5):
                temp_line += line[i + j]
            i += 5
            if len(temp_line) == 5:
                line_by_five.append(temp_line)

        except:
            break

    decoded_five_line = code_baudot(line_by_five)
    print(f'Baudot decoding: {decoded_five_line}')

    byte_line = bytearray(int(line[i:i + 8], 2) for i in range(0, len(line), 8))

    print(f"KOI-8R decoding: {byte_line.decode('koi8-r')}")
    print(f"cp866 decoding: {byte_line.decode('cp866')}")
    print(f"windows1251: {byte_line.decode('windows-1251')}")

path_d = 'variant12.docx'
print(f'doc is {path_d}\n')

parameter_d, flags_d = find_parameter(path_d)
print(f'parameter - {parameter_d}\n')

line_d = zeros_ones(path_d, parameter_d, flags_d)
print(line_d)

decode_d_array = decode(line_d)